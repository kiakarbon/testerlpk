from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime
import tempfile

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Borders
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    # List of all border tags
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            
            # Check for existing tag
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), edge_data[key])

def create_word_document(catatan):
    """
    Membuat dokumen Word profesional dari catatan praktik
    """
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Judul utama
    title = doc.add_heading('LAPORAN PRAKTIKUM NANOMATERIAL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    title.runs[0].font.size = Pt(16)
    title.runs[0].bold = True
    
    # Sub judul
    subtitle = doc.add_heading(catatan['judul'], 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()  # Spasi
    
    # Informasi metadata dalam tabel
    table_info = doc.add_table(rows=6, cols=2)
    table_info.style = 'Light Shading'
    
    data_info = [
        ("Nama Praktikan", catatan['nama_praktikan']),
        ("Tanggal Praktikum", catatan['tanggal']),
        ("Institusi/Laboratorium", catatan.get('institusi', '-')),
        ("Kelompok/Shift", catatan.get('kelompok', '-')),
        ("Supervisor", catatan.get('supervisor', '-')),
        ("Timestamp", catatan['timestamp'])
    ]
    
    for i, (label, value) in enumerate(data_info):
        table_info.cell(i, 0).text = label
        table_info.cell(i, 1).text = value
        table_info.cell(i, 0).paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()  # Spasi
    
    # Spesifikasi Nanomaterial
    doc.add_heading('SPESIFIKASI NANOMATERIAL', level=2)
    table_nano = doc.add_table(rows=3, cols=2)
    table_nano.style = 'Light Shading'
    
    nano_data = [
        ("Jenis Nanomaterial", catatan['jenis_nanomaterial']),
        ("Metode Sintesis", catatan['metode_sintesis']),
        ("Pelarut", catatan['pelarut'])
    ]
    
    for i, (label, value) in enumerate(nano_data):
        table_nano.cell(i, 0).text = label
        table_nano.cell(i, 1).text = value
        table_nano.cell(i, 0).paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()  # Spasi
    
    # Parameter Sintesis
    doc.add_heading('PARAMETER SINTESIS', level=2)
    table_param = doc.add_table(rows=6, cols=2)
    table_param.style = 'Light Shading'
    
    param_data = [
        ("Suhu Sintesis", f"{catatan['suhu']} °C"),
        ("Waktu Sintesis", f"{catatan['waktu']} jam"),
        ("Tekanan", f"{catatan.get('tekanan', '-')} atm"),
        ("pH Larutan", f"{catatan['ph']}"),
        ("Konsentrasi", f"{catatan['konsentrasi']} mg/mL"),
        ("Total Parameter", "6 parameter tercatat")
    ]
    
    for i, (label, value) in enumerate(param_data):
        table_param.cell(i, 0).text = label
        table_param.cell(i, 1).text = value
        table_param.cell(i, 0).paragraphs[0].runs[0].bold = True
    
    doc.add_page_break()
    
    # Prosedur Praktik
    doc.add_heading('PROSEDUR PRAKTIKUM', level=2)
    prosedur_paragraph = doc.add_paragraph()
    prosedur_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    for line in catatan['prosedur'].split('\n'):
        if line.strip():
            prosedur_paragraph.add_run(line.strip() + '\n')
    
    doc.add_paragraph()  # Spasi
    
    # Hasil Pengamatan
    doc.add_heading('HASIL PENGAMATAN', level=2)
    hasil_paragraph = doc.add_paragraph()
    hasil_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    for line in catatan['hasil_pengamatan'].split('\n'):
        if line.strip():
            hasil_paragraph.add_run(line.strip() + '\n')
    
    # Gambar hasil jika ada
    if catatan.get('image_paths'):
        doc.add_paragraph()  # Spasi
        doc.add_heading('DOKUMENTASI VISUAL', level=2)
        
        for img_path in catatan['image_paths'][:3]:  # Maks 3 gambar
            try:
                doc.add_picture(img_path, width=Inches(5))
                doc.add_paragraph(f"Gambar: {os.path.basename(img_path)}")
                doc.add_paragraph()  # Spasi
            except:
                pass
    
    # Catatan Tambahan
    if catatan.get('catatan_tambahan'):
        doc.add_heading('CATATAN TAMBAHAN', level=2)
        catatan_paragraph = doc.add_paragraph()
        catatan_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        catatan_paragraph.add_run(catatan['catatan_tambahan'])
    
    # Footer dengan informasi pembuatan
    doc.add_paragraph()  # Spasi
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f"Dokumen dibuat otomatis oleh Lab PSA Nano • {datetime.now().strftime('%d %B %Y %H:%M:%S')}")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_run.italic = True
    
    # Simpan ke temporary file
    temp_dir = tempfile.gettempdir()
    safe_title = "".join(c for c in catatan['judul'] if c.isalnum() or c in (' ', '-', '_')).rstrip()
    filename = f"Laporan_{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(temp_dir, filename)
    doc.save(filepath)
    
    return filepath

def create_batch_word_document(catatan_list):
    """
    Membuat dokumen Word gabungan untuk multiple catatan
    """
    doc = Document()
    
    # Judul batch
    title = doc.add_heading('KOMPILASI LAPORAN PRAKTIKUM NANOMATERIAL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    
    doc.add_paragraph(f"Total Laporan: {len(catatan_list)}")
    doc.add_paragraph(f"Tanggal Kompilasi: {datetime.now().strftime('%d %B %Y')}")
    
    doc.add_page_break()
    
    # Tambahkan setiap catatan
    for idx, catatan in enumerate(catatan_list, 1):
        doc.add_heading(f'LAPORAN {idx}: {catatan["judul"]}', level=1)
        
        # Informasi dasar
        doc.add_paragraph(f"Praktikan: {catatan['nama_praktikan']}")
        doc.add_paragraph(f"Tanggal: {catatan['tanggal']}")
        doc.add_paragraph(f"Nanomaterial: {catatan['jenis_nanomaterial']}")
        doc.add_paragraph(f"Metode: {catatan['metode_sintesis']}")
        
        # Prosedur singkat
        doc.add_heading('Prosedur Singkat', level=2)
        first_lines = '\n'.join(catatan['prosedur'].split('\n')[:3])
        doc.add_paragraph(first_lines + '...')
        
        doc.add_page_break()
    
    # Simpan file
    temp_dir = tempfile.gettempdir()
    filename = f"Batch_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(temp_dir, filename)
    doc.save(filepath)
    
    return filepath
